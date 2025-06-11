import io
import pytest
from fastapi.testclient import TestClient
from app.main import app
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook

# -----------------------------
# Helper Functions
# -----------------------------
def create_test_pdf() -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(100, 750, "Hello PDF")
    c.save()
    buffer.seek(0)
    return buffer.read()

def create_test_docx() -> bytes:
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("Test DOCX", 0)
    doc.add_paragraph("This is a test document.")
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def create_test_xlsx() -> bytes:
    buffer = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Age"])
    ws.append(["Alice", 30])
    ws.append(["Bob", 25])
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generate_dummy_file(size_bytes: int) -> io.BytesIO:
    return io.BytesIO(b"a" * size_bytes)


client = TestClient(app)

# -----------------------------
# Test Data Setup
# -----------------------------

VALID_TEST_FILES = {
    "txt": (b"Hello, this is a text file.", "text/plain"),
    "pdf": (create_test_pdf(), "application/pdf"),
    "csv": (b"name,age\nAlice,30\nBob,25", "text/csv"),
    "json": (b'{"name": "Alice", "age": 30}', "application/json"),
    "xml": (b"<person><name>Alice</name><age>30</age></person>", "application/xml"),
    "html": (b"<html><body><h1>Hello</h1></body></html>", "text/html"),
    "docx": (create_test_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "xlsx": (create_test_xlsx(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
}

MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE = MAX_FILE_SIZE_MB * 1024 * 1024

# -----------------------------
# Valid File Tests
# -----------------------------

@pytest.mark.parametrize("ext, file_info", VALID_TEST_FILES.items())
def test_valid_file_upload_returns_200(ext, file_info):
    """
    Ensure valid file types and content return 200 OK with proper metadata.
    """
    content, content_type = file_info
    filename = f"testfile.{ext}"
    file = {"file": (filename, io.BytesIO(content), content_type)}

    response = client.post("/file", files=file)

    assert response.status_code == 200
    json_data = response.json()
    assert "text" in json_data
    assert json_data["metadata"]["filename"] == filename

# -----------------------------
# Invalid File Type Test
# -----------------------------

def test_invalid_extension_and_content_type_returns_400():
    """
    Uploading a disallowed file extension and MIME type should return 400.
    """
    file = {
        "file": ("testfile.exe", io.BytesIO(b"dummy data"), "application/octet-stream")
    }

    response = client.post("/file", files=file)

    assert response.status_code == 400
    assert response.json()["detail"] == "File type not allowed"

# -----------------------------
# File Size Tests
# -----------------------------

def test_file_within_max_size_returns_200():
    """
    File just under max size limit should be accepted.
    """
    file = {
        "file": ("small.txt", generate_dummy_file(MAX_FILE_SIZE - 1), "text/plain")
    }

    response = client.post("/file", files=file)

    assert response.status_code == 200
    assert response.json()["metadata"]["filename"] == "small.txt"

def test_file_exceeding_max_size_returns_413():
    """
    File over max size limit should be rejected with 413.
    """
    file = {
        "file": ("large.txt", generate_dummy_file(MAX_FILE_SIZE + 1), "text/plain")
    }

    response = client.post("/file", files=file)

    assert response.status_code == 413
    assert "File too large" in response.json()["detail"]
